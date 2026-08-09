#!/usr/bin/env python3
"""
Render an MB-branded .docx from a simple JSON report structure.
Used by the SEO/GEO skills to export any deliverable (audit, content draft,
gap matrix, competitor scorecard, backlink report, GEO visibility report) as
a docx with the MB logo in the page header, instead of leaving it in chat.

Usage: python3 render_mb_docx.py <input.json> <output.docx>

Input JSON shape:
{
  "title": "Technical SEO Audit",
  "subtitle": "https://www.mauriceblackburn.com.au/... (optional URL/scope)",
  "date": "2026-07-12",
  "sections": [
    {"heading": "Summary", "body": ["one or more paragraphs"]},
    {"heading": "Findings", "table": {"headers": [...], "rows": [[...], ...]}},
    {"heading": "Action plan", "list": ["item 1", "item 2"]},
    {"heading": "Schema", "code": "raw text or JSON, monospaced block"},
    {"heading": "Page Copy: Section Title", "body": [...], "subsections": [
      {"heading": "An H3 sub-heading within this H2 section", "body": [...]}
    ]}
  ]
}

Only "title" and "sections" are required. Each section needs "heading" plus
any of body/table/list/code, and optionally "subsections" (a list of the
same shape, one nesting level deeper, rendered as a real Word Heading 2
under the section's Heading 1). Use "subsections" to carry a drafted
page's actual H2/H3 structure into the docx as real heading styles, not
flat body text, so the hierarchy is visible in Word's Navigation pane.

This file is the deliverable itself: never write an AI/assistant
self-reference, a "reporter" or "prepared by AI" byline, or first-person
orchestrator commentary into the title, sections, or date fields. It must
read as ready for internal use as-is.
"""
import json
import sys
from pathlib import Path

BRAND_PRIMARY = "461C2F"   # MB Shade 4, primary heading/dark UI
BRAND_ACCENT = "DB333C"    # MB Red
BRAND_MUTED = "4E4E4E"     # Dark Grey, body/meta text
LOGO_PATH = Path(__file__).resolve().parents[2] / "brand-mb" / "assets" / "mb-logo-docx-header.png"


def add_header_logo(doc):
    from docx.shared import Inches
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    run = p.add_run()
    if LOGO_PATH.exists():
        run.add_picture(str(LOGO_PATH), width=Inches(1.8))
    else:
        run.text = "Maurice Blackburn"


def add_title_block(doc, title, subtitle, date):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    h = doc.add_heading(level=0)
    run = h.add_run(title)
    run.font.color.rgb = RGBColor.from_string(BRAND_PRIMARY)

    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(BRAND_MUTED)

    if date:
        p = doc.add_paragraph()
        run = p.add_run(f"Last updated {date}")
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor.from_string(BRAND_MUTED)

    # thin brand-red rule under the title block
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BRAND_ACCENT)
    pbdr.append(bottom)
    pPr.append(pbdr)


import re as _re

def _add_hyperlink(paragraph, url, text):
    """Add a real clickable hyperlink run to a python-docx paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "CC2E57"); rPr.append(color)  # MB Shade 1, link
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_rich_text(paragraph, text):
    """Render text containing [anchor](url) markdown links as real runs + hyperlinks."""
    pos = 0
    for m in _re.finditer(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        _add_hyperlink(paragraph, m.group(2), m.group(1))
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_section_body(doc, section):
    from docx.shared import Pt

    if "body" in section:
        for para_text in section["body"]:
            p = doc.add_paragraph()
            _add_rich_text(p, para_text)

    if "list" in section:
        for item in section["list"]:
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, item)

    if "table" in section:
        headers = section["table"]["headers"]
        rows = section["table"]["rows"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, htext in enumerate(headers):
            hdr_cells[i].text = str(htext)
            for run in hdr_cells[i].paragraphs[0].runs:
                run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

    if "code" in section:
        p = doc.add_paragraph()
        run = p.add_run(section["code"])
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_section(doc, section, level=1):
    from docx.shared import RGBColor, Pt

    # Clear, distinct visual hierarchy per level (users found the default H1/H2/H3 too similar):
    # size, colour, and space-before all step down by level.
    level_style = {
        1: {"size": 20, "color": BRAND_PRIMARY, "space_before": 18},   # page H1 / top section
        2: {"size": 15, "color": BRAND_PRIMARY, "space_before": 14},   # page H2 / section
        3: {"size": 12.5, "color": BRAND_ACCENT, "space_before": 10},  # page H3 / FAQ question
    }[min(level, 3)]

    h = doc.add_heading(level=min(level, 3))
    run = h.add_run(section["heading"])
    run.font.color.rgb = RGBColor.from_string(level_style["color"])
    run.font.size = Pt(level_style["size"])
    run.bold = True
    h.paragraph_format.space_before = Pt(level_style["space_before"])
    h.paragraph_format.space_after = Pt(4)

    add_section_body(doc, section)

    for sub in section.get("subsections", []):
        add_section(doc, sub, level=min(level + 1, 3))


def render(data, out_path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_header_logo(doc)
    add_title_block(doc, data["title"], data.get("subtitle"), data.get("date"))

    for section in data.get("sections", []):
        add_section(doc, section)

    doc.save(out_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 render_mb_docx.py <input.json> <output.docx>", file=sys.stderr)
        sys.exit(1)
    input_path, out_path = sys.argv[1], sys.argv[2]
    data = json.loads(Path(input_path).read_text())
    render(data, out_path)
    print(f"DOCX -> {out_path}")


if __name__ == "__main__":
    main()
